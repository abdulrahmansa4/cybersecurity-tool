import ttkbootstrap as tb
from ttkbootstrap.constants import *
from tkinter.scrolledtext import ScrolledText
from ttkbootstrap.dialogs import Messagebox

import subprocess
import re
import queue
import threading
import json
import csv
from datetime import datetime
import os
import winsound
import requests
import matplotlib
matplotlib.use("TkAgg")
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import time
import logging
import socket
import shutil
import geoip2.database
from tkinter import filedialog
import platform


logging.basicConfig(filename='analyzer.log', level=logging.INFO)

class WireSharkAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("WireShark Security Analyzer")
        self.root.geometry("1280x850")
        self.nmap_process = None
        self.last_sound_time = 0

        # Capture control
        self.is_capturing = False
        self.capture_process = None
        self.interface = None

        # Packet storage
        self.packets = []
        self.threats = []

        # AI model
        self.model_name = "deepseek-coder"

        # UI Setup
        self.setup_ui()

        # Get interfaces
        self.refresh_interfaces()

        # Setup statistics tab
        self.setup_stats_tab()

        # Dark mode flag
        self.dark_mode = False

        self.packet_queue = queue.Queue()
        self.root.after(100, self.process_packet_queue)

        

        # Status bar
        self.status_var = tb.StringVar()
        status_bar = tb.Label(self.root, textvariable=self.status_var, relief=tb.SUNKEN, anchor="w")
        status_bar.pack(side=tb.BOTTOM, fill=tb.X)
        self.status_var.set("Ready")

        # Bindings
        self.root.bind('<Control-s>', lambda e: self.export_data())

        # Open GeoLite2 database once (always relative to script location)
        try:
            db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "GeoLite2-Country.mmdb")
            self.geoip_reader = geoip2.database.Reader(db_path)
        except Exception as e:
            self.geoip_reader = None
            Messagebox.show_error("GeoIP Error", f"Could not open GeoLite2-Country.mmdb: {e}")

    def get_country_flag(self, ip):
        private_ranges = [
            re.compile(r"^10\."), 
            re.compile(r"^192\.168\."), 
            re.compile(r"^172\.(1[6-9]|2[0-9]|3[0-1])\.")
        ]
        if any(r.match(ip) for r in private_ranges):
            return "Private"
        if not self.geoip_reader:
            return "Unknown"
        try:
            response = self.geoip_reader.country(ip)
            country = response.country.name or "Unknown"
            code = response.country.iso_code or ""
            flag = ""
            if code and len(code) == 2:
                flag = chr(127397 + ord(code[0].upper())) + chr(127397 + ord(code[1].upper()))
            return f"{country} {flag}".strip()
        except Exception:
            return "Unknown"

    def setup_ui(self):



        # Main container
        main_frame = tb.Frame(self.root, padding=10)
        main_frame.pack(fill=BOTH, expand=True)

        # Tabbed notebook
        self.notebook = tb.Notebook(main_frame)
        self.notebook.pack(fill=BOTH, expand=True)

        # Tabs
        self.tab_packets = tb.Frame(self.notebook)
        self.tab_threats = tb.Frame(self.notebook)
        self.tab_stats = tb.Frame(self.notebook)
        self.tab_settings = tb.Frame(self.notebook)
        self.tab_chat = tb.Frame(self.notebook)
        self.tab_nmap = tb.Frame(self.notebook)
        self.tab_pentest = tb.Frame(self.notebook)
        self.notebook.add(self.tab_packets, text="Live Packets")
        self.notebook.add(self.tab_threats, text="Threats")
        self.notebook.add(self.tab_stats, text="Statistics")
        self.notebook.add(self.tab_chat, text="AI Chat")
        self.notebook.add(self.tab_nmap, text="Nmap Analyzer")
        self.notebook.add(self.tab_settings, text="Settings")
        self.notebook.add(self.tab_pentest, text="Pentest Tools")

        # --- Live Packets Tab ---
        control_frame = tb.Labelframe(self.tab_packets, text="Capture Controls", padding=10)
        control_frame.pack(fill=X, pady=(0, 10), padx=10)

        tb.Label(control_frame, text="Interface:").grid(row=0, column=0, padx=5)
        self.interface_combo = tb.Combobox(control_frame, width=30)
        self.interface_combo.grid(row=0, column=1, padx=5)
        tb.Button(control_frame, text="Refresh", command=self.refresh_interfaces, bootstyle="info-outline").grid(row=0, column=2, padx=5)

        tb.Label(control_frame, text="Filter:").grid(row=0, column=3, padx=5)
        self.filter_entry = tb.Entry(control_frame, width=30)
        self.filter_entry.grid(row=0, column=4, padx=5)
        self.filter_entry.insert(0, "tcp or udp")

        self.start_btn = tb.Button(control_frame, text="Start", command=self.start_capture, bootstyle="success")
        self.start_btn.grid(row=0, column=5, padx=5)
        tb.Button(control_frame, text="Stop", command=self.stop_capture, bootstyle="danger").grid(row=0, column=6, padx=5)
        tb.Button(control_frame, text="Clear", command=self.clear_data, bootstyle="warning-outline").grid(row=0, column=7, padx=5)

        # Export/Report Buttons
        export_frame = tb.Frame(control_frame)
        export_frame.grid(row=0, column=8, columnspan=6, padx=(20,0))
        tb.Button(export_frame, text="Export", command=self.export_packets_prompt, bootstyle="secondary").pack(side=LEFT, padx=2)
       

        # Search
        tb.Label(control_frame, text="Search:").grid(row=0, column=14, padx=5)
        self.search_entry = tb.Entry(control_frame, width=20)
        self.search_entry.grid(row=0, column=15, padx=5)
        tb.Button(control_frame, text="Go", command=self.search_packets, bootstyle="info-outline").grid(row=0, column=16, padx=5)

        # Main content area
        content_frame = tb.Frame(self.tab_packets)
        content_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # Packet display
        packet_frame = tb.Labelframe(content_frame, text="Packet Stream", padding=10)
        packet_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 10))

        self.packet_tree = tb.Treeview(
            packet_frame,
            columns=("No", "Time", "Source", "Src Country", "Destination", "Dst Country", "Protocol", "Length", "Info"),
            show="headings"
        )
        for col, width in zip(
            ("No", "Time", "Source", "Src Country", "Destination", "Dst Country", "Protocol", "Length", "Info"),
            (50, 120, 150, 120, 150, 120, 80, 80, 300)
        ):
            self.packet_tree.heading(col, text=col)
            self.packet_tree.column(col, width=width)
        self.packet_tree.pack(fill=BOTH, expand=True)
        self.packet_tree.bind("<Double-1>", self.show_packet_details)

        # --- Threats Tab ---
        threat_tree_frame = tb.Labelframe(self.tab_threats, text="Detected Threats", padding=10)
        threat_tree_frame.pack(fill=BOTH, expand=True, padx=10, pady=(10, 0))

        self.threat_tree = tb.Treeview(
            threat_tree_frame,
            columns=("No", "Time", "Source", "Src Country", "Destination", "Dst Country", "Threat", "Info", "AI Explain", "Fix"),
            show="headings"
        )
        for col, width in zip(
            ("No", "Time", "Source", "Src Country", "Destination", "Dst Country", "Threat", "Info", "AI Explain", "Fix"),
            (50, 120, 150, 120, 150, 120, 120, 200, 200, 200)
        ):
            self.threat_tree.heading(col, text=col)
            self.threat_tree.column(col, width=width)
        self.threat_tree.pack(fill=BOTH, expand=True)
        self.threat_tree.bind("<Double-1>", self.show_threat_details)

        # Threat summary and stats
        summary_stats_frame = tb.Frame(self.tab_threats)
        summary_stats_frame.pack(fill=X, padx=10, pady=(10, 0))

        self.threat_summary = tb.Text(summary_stats_frame, height=5, state="disabled", font=("Segoe UI", 11))
        self.threat_summary.pack(fill=X, side=LEFT, expand=True, padx=(0, 10))

        stats_frame = tb.Frame(summary_stats_frame)
        stats_frame.pack(side=RIGHT, fill=Y)
        tb.Label(stats_frame, text="Packets:").pack(side=LEFT)
        self.packet_count = tb.Label(stats_frame, text="0")
        self.packet_count.pack(side=LEFT, padx=(0, 10))
        tb.Label(stats_frame, text="Threats:").pack(side=LEFT)
        self.threat_count = tb.Label(stats_frame, text="0")
        self.threat_count.pack(side=LEFT)

        # Threats tab buttons
        threat_btn_frame = tb.Frame(self.tab_threats)
        threat_btn_frame.pack(fill=X, padx=10, pady=(10, 10))
        tb.Button(threat_btn_frame, text="Analyze Threats with AI", command=self.analyze_threats_with_ai, bootstyle="info").pack(side=LEFT, padx=2)
        tb.Button(threat_btn_frame, text="Create Threats Report", command=self.create_ai_report, bootstyle="secondary").pack(side=LEFT, padx=2)
        tb.Button(threat_btn_frame, text="Export Threats CSV", command=self.export_threats_csv, bootstyle="secondary").pack(side=LEFT, padx=2)
        tb.Button(threat_btn_frame, text="Export Threats PDF", command=self.export_threats_pdf, bootstyle="secondary").pack(side=LEFT, padx=2)

        

        # --- Add Color Legend at the bottom of Threats Tab ---
        legend_frame = tb.Frame(self.tab_threats)
        legend_frame.pack(side="bottom", fill=X, padx=10, pady=5)

        tb.Label(legend_frame, text="Color Legend: ").pack(side=LEFT)
        tb.Label(legend_frame, text="Malware", foreground="#FF1744").pack(side=LEFT, padx=5)
        tb.Label(legend_frame, text="DDoS", foreground="#FF9100").pack(side=LEFT, padx=5)
        tb.Label(legend_frame, text="Port Scan", foreground="#2979FF").pack(side=LEFT, padx=5)
        tb.Label(legend_frame, text="IP Exposure", foreground="#D500F9").pack(side=LEFT, padx=5)
        tb.Label(legend_frame, text="Error/Fail", foreground="#C51162").pack(side=LEFT, padx=5)
        tb.Label(legend_frame, text="Other", foreground="#222222").pack(side=LEFT, padx=5)

        # --- Statistics Tab ---
        # (setup_stats_tab is called in __init__)

        # --- Settings Tab ---
        settings_frame = tb.Frame(self.tab_settings, padding=10)
        settings_frame.pack(fill=BOTH, expand=True)

        tb.Label(settings_frame, text="AI Model:").pack(anchor="w", padx=10, pady=5)
        self.model_entry = tb.Entry(settings_frame)
        self.model_entry.insert(0, "deepseek-coder")
        self.model_entry.pack(anchor="w", padx=10)

        # Advanced user options
        self.sound_var = tb.BooleanVar(value=True)
        self.autoscroll_var = tb.BooleanVar(value=True)
        self.darkmode_var = tb.BooleanVar(value=False)
        self.autosave_var = tb.BooleanVar(value=False)
        self.show_tooltips_var = tb.BooleanVar(value=True)
        self.confirm_stop_var = tb.BooleanVar(value=True)
        self.save_logs_var = tb.BooleanVar(value=True)
        self.show_ai_explain_var = tb.BooleanVar(value=True)
        self.show_ai_fix_var = tb.BooleanVar(value=True)

        tb.Checkbutton(settings_frame, text="Enable Notification Sound", variable=self.sound_var, bootstyle="success").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Auto-scroll Packet/Threat Views", variable=self.autoscroll_var, bootstyle="info").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Dark Mode", variable=self.darkmode_var, command=self.toggle_dark_mode, bootstyle="dark").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Auto-save Captures", variable=self.autosave_var, bootstyle="secondary").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Show Tooltips", variable=self.show_tooltips_var, bootstyle="info").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Confirm Before Stopping Scan", variable=self.confirm_stop_var, bootstyle="warning").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Save Logs to File", variable=self.save_logs_var, bootstyle="secondary").pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Show AI Explanation in Threats", variable=self.show_ai_explain_var, bootstyle="info", command=self.update_threat_analysis).pack(anchor="w", padx=10, pady=2)
        tb.Checkbutton(settings_frame, text="Show AI Fix in Threats", variable=self.show_ai_fix_var, bootstyle="info", command=self.update_threat_analysis).pack(anchor="w", padx=10, pady=2)

        tb.Button(settings_frame, text="Save Settings", command=self.save_settings, bootstyle="success").pack(anchor="w", padx=10, pady=10)

        # --- AI Chat Tab ---
        chat_frame = tb.Frame(self.tab_chat, padding=10)
        chat_frame.pack(fill=BOTH, expand=True)

        chat_label = tb.Label(chat_frame, text="AI Security Assistant", font=("Segoe UI", 14, "bold"))
        chat_label.pack(anchor="w", pady=(0, 5))

        self.chat_display = ScrolledText(chat_frame, height=20, state="disabled", wrap="word", font=("Segoe UI", 13))
        self.chat_display.pack(fill=BOTH, expand=True, pady=(0, 10))

        input_frame = tb.Frame(chat_frame)
        input_frame.pack(fill=X)
        self.chat_entry = tb.Entry(input_frame, font=("Segoe UI", 12))
        self.chat_entry.pack(side=LEFT, fill=X, expand=True, padx=(0, 5))
        send_btn = tb.Button(input_frame, text="Send", command=self.send_chat_message, bootstyle="primary")
        send_btn.pack(side=LEFT)

        # Chat message tags
        self.chat_display.tag_config("user", foreground="#007acc", font=("Segoe UI", 12, "bold"))
        self.chat_display.tag_config("ai", foreground="#FF8800", font=("Segoe UI", 13, "italic"))
        self.chat_display.tag_config("thinking", foreground="#888888", font=("Segoe UI", 12, "italic"))

        self.chat_entry.bind("<Return>", lambda event: self.send_chat_message())

        # --- Nmap Analyzer Tab ---
        nmap_frame = tb.Frame(self.tab_nmap, padding=10)
        nmap_frame.pack(fill=BOTH, expand=True)

        # Top: Target and Scan Type
        nmap_target_frame = tb.Labelframe(nmap_frame, text="Target & Scan Type", padding=10)
        nmap_target_frame.pack(fill=tb.X, pady=(0, 10))
        tb.Label(nmap_target_frame, text="Target IP/Domain:").grid(row=0, column=0, sticky="w")
        self.nmap_target_entry = tb.Entry(nmap_target_frame, width=30)
        self.nmap_target_entry.grid(row=0, column=1, padx=5)
        tb.Label(nmap_target_frame, text="Scan Type:").grid(row=0, column=2, padx=5)
        self.nmap_scan_type = tb.Combobox(nmap_target_frame, width=12, state="readonly")
        self.nmap_scan_type["values"] = ["Fast", "Medium", "Full", "Custom"]
        self.nmap_scan_type.current(0)
        self.nmap_scan_type.grid(row=0, column=3, padx=5)
        tb.Button(nmap_target_frame, text="Run Nmap Scan", command=self.run_nmap_scan).grid(row=0, column=4, padx=5)
        self.nmap_stop_btn = tb.Button(nmap_target_frame, text="Stop Scan", command=self.stop_nmap_scan, state="disabled")
        self.nmap_stop_btn.grid(row=0, column=5, padx=5)

        # Middle: Scan Options
        nmap_options_frame = tb.Labelframe(nmap_frame, text="Scan Options", padding=10)
        nmap_options_frame.pack(fill=tb.X, pady=(0, 10))
        tb.Label(nmap_options_frame, text="Ports:").grid(row=0, column=0, padx=5, sticky="w")
        self.nmap_ports_entry = tb.Entry(nmap_options_frame, width=15)
        self.nmap_ports_entry.grid(row=0, column=1, padx=5)
        self.nmap_ports_entry.insert(0, "")

        tb.Label(nmap_options_frame, text="Timing:").grid(row=0, column=2, padx=5, sticky="w")
        self.nmap_timing_combo = tb.Combobox(nmap_options_frame, width=5, state="readonly")
        self.nmap_timing_combo["values"] = ["T0", "T1", "T2", "T3", "T4", "T5"]
        self.nmap_timing_combo.current(4)
        self.nmap_timing_combo.grid(row=0, column=3, padx=5)

        tb.Label(nmap_options_frame, text="NSE Scripts:").grid(row=0, column=4, padx=5, sticky="w")
        self.nmap_scripts_entry = tb.Entry(nmap_options_frame, width=25)
        self.nmap_scripts_entry.grid(row=0, column=5, padx=5)
        self.nmap_scripts_entry.insert(0, "")

        # Checkboxes: Aggressive, OS, Service
        self.nmap_aggressive_var = tb.BooleanVar()
        self.nmap_osdetect_var = tb.BooleanVar()
        self.nmap_service_var = tb.BooleanVar()
        tb.Checkbutton(nmap_options_frame, text="Aggressive (-A)", variable=self.nmap_aggressive_var).grid(row=1, column=0, padx=5, sticky="w")
        tb.Checkbutton(nmap_options_frame, text="OS Detect (-O)", variable=self.nmap_osdetect_var).grid(row=1, column=1, padx=5, sticky="w")
        tb.Checkbutton(nmap_options_frame, text="Service Version (-sV)", variable=self.nmap_service_var).grid(row=1, column=2, padx=5, sticky="w")

        # Checkboxes: IPv6, No Ping, Verbose, Debug
        self.nmap_ipv6_var = tb.BooleanVar()
        self.nmap_no_ping_var = tb.BooleanVar()
        self.nmap_verbose_var = tb.BooleanVar()
        self.nmap_debug_var = tb.BooleanVar()
        tb.Checkbutton(nmap_options_frame, text="IPv6 (-6)", variable=self.nmap_ipv6_var).grid(row=1, column=3, padx=5, sticky="w")
        tb.Checkbutton(nmap_options_frame, text="No Ping (-Pn)", variable=self.nmap_no_ping_var).grid(row=1, column=4, padx=5, sticky="w")
        tb.Checkbutton(nmap_options_frame, text="Verbose (-v)", variable=self.nmap_verbose_var).grid(row=1, column=5, padx=5, sticky="w")
        tb.Checkbutton(nmap_options_frame, text="Debug (-d)", variable=self.nmap_debug_var).grid(row=1, column=6, padx=5, sticky="w")

        # Output and Extra Args
        nmap_output_frame = tb.Labelframe(nmap_frame, text="Output & Profiles", padding=10)
        nmap_output_frame.pack(fill=tb.X, pady=(0, 10))
        tb.Label(nmap_output_frame, text="Output:").grid(row=0, column=0, padx=5, sticky="w")
        self.nmap_output_format = tb.Combobox(nmap_output_frame, width=10, state="readonly")
        self.nmap_output_format["values"] = ["Normal", "XML", "Grepable"]
        self.nmap_output_format.current(0)
        self.nmap_output_format.grid(row=0, column=1, padx=5)
        tb.Label(nmap_output_frame, text="Extra Args:").grid(row=0, column=2, padx=5, sticky="w")
        self.nmap_extra_entry = tb.Entry(nmap_output_frame, width=40)
        self.nmap_extra_entry.grid(row=0, column=3, padx=5)
        tb.Button(nmap_output_frame, text="Save Profile", command=self.save_nmap_profile).grid(row=0, column=4, padx=5)
        tb.Button(nmap_output_frame, text="Load Profile", command=self.load_nmap_profile).grid(row=0, column=5, padx=5)

        # Output area
        self.nmap_output = ScrolledText(nmap_frame, height=20, state="normal", wrap="word", font=("Segoe UI", 12))
        self.nmap_output.pack(fill=BOTH, expand=True, pady=(10, 10))

        nmap_btn_frame = tb.Frame(nmap_frame)
        nmap_btn_frame.pack(fill=tb.X, pady=(0, 10))
        tb.Button(nmap_btn_frame, text="Analyze with AI", command=self.analyze_nmap_with_ai).pack(side=LEFT, padx=2)
        tb.Button(nmap_btn_frame, text="Create Nmap Report", command=self.create_nmap_report).pack(side=LEFT, padx=2)

        self.nmap_progress = tb.Progressbar(nmap_frame, mode="determinate", maximum=100)
        self.nmap_progress.pack(fill=tb.X, pady=(5, 0))
        self.nmap_progress_label = tb.Label(nmap_frame, text="", font=("Segoe UI", 11))
        self.nmap_progress_label.pack(fill=tb.X)
        self.nmap_progress_label.pack_forget()

        style = tb.Style()
        style.configure("Treeview", font=("Segoe UI Emoji", 10))
                # ...inside setup_ui(), after Nmap Analyzer tab setup...
        
        # --- Pentest Tools Tab ---
        pentest_frame = tb.Frame(self.tab_pentest, padding=10)
        pentest_frame.pack(fill=BOTH, expand=True)
        
        # Target entry
        tb.Label(pentest_frame, text="Target:").pack(anchor="w")
        self.pentest_target_entry = tb.Entry(pentest_frame, width=40)
        self.pentest_target_entry.pack(anchor="w", pady=(0, 10))
        
        # Buttons for tools
        btn_frame = tb.Frame(pentest_frame)
        btn_frame.pack(anchor="w", pady=(0, 10))
        tb.Button(btn_frame, text="Ping", command=self.run_ping).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="Traceroute", command=self.run_traceroute).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="Whois", command=self.run_whois).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="Nslookup", command=self.run_nslookup).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="Port Scan", command=self.run_portscan).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="HTTP Headers", command=self.run_http_headers).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="SSL Info", command=self.run_ssl_info).pack(side=LEFT, padx=2)
        tb.Button(btn_frame, text="DNS Zone Transfer", command=self.run_dns_zone_transfer).pack(side=LEFT, padx=2)
        
        # Output area
        self.pentest_output = ScrolledText(pentest_frame, height=18, state="normal", wrap="word", font=("Segoe UI", 12))
        self.pentest_output.pack(fill=BOTH, expand=True, pady=(10, 0))
        self.pentest_output.tag_config("header", foreground="#FF8800", font=("Segoe UI", 12, "bold"))
        self.pentest_output.tag_config("output", foreground="#FFFFFF", font=("Segoe UI", 11))

    def setup_stats_tab(self):
        """Setup the statistics tab with a matplotlib chart"""
        fig = Figure(figsize=(5, 2), dpi=100)
        self.ax = fig.add_subplot(111)
        self.ax.set_title("Threats Over Time")
        self.ax.set_xlabel("Packet #")
        self.ax.set_ylabel("Threat Count")
        self.stats_canvas = FigureCanvasTkAgg(fig, master=self.tab_stats)
        self.stats_canvas.get_tk_widget().pack(fill=tb.BOTH, expand=True)
        self.stats_data = []

    def refresh_interfaces(self):
        """Get available network interfaces using tshark"""
        try:
            result = subprocess.run(
                ["C:\\Program Files\\Wireshark\\tshark.exe", "-D"],
                capture_output=True,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW
            )
            if result.returncode != 0:
                Messagebox.show_error("Error", "Could not list interfaces. Is Wireshark installed?")
                return
            interfaces = []
            self.interface_map = {}
            for line in result.stdout.splitlines():
                if match := re.match(r"\d+\. ([^ ]+) \((.+)\)", line):
                    device = match.group(1)
                    friendly = match.group(2)
                    display = f"{friendly} [{device}]"
                    interfaces.append(display)
                    self.interface_map[display] = device
            self.interface_combo["values"] = interfaces
            if interfaces:
                self.interface_combo.current(0)
        except Exception as e:
            Messagebox.show_error("Error", f"Failed to get interfaces: {str(e)}")

    def start_capture(self):
        """Start packet capture"""
        if self.is_capturing:
            return
        self.interface = self.interface_combo.get()
        if not self.interface:
            Messagebox.show_error("Error", "Please select a network interface")
            return
        self.interface = self.interface_map[self.interface]
        self.is_capturing = True
        self.start_btn.config(text="Capturing...")
        capture_thread = threading.Thread(
            target=self.run_tshark_capture,
            daemon=True
        )
        capture_thread.start()

    def run_tshark_capture(self):
        """Run tshark capture process"""
        try:
            filter_str = self.filter_entry.get()
            self.capture_process = subprocess.Popen(
                [
                    "C:\\Program Files\\Wireshark\\tshark.exe",
                    "-i", self.interface,
                    "-Y", filter_str,
                    "-T", "fields",
                    "-e", "frame.number",
                    "-e", "frame.time",
                    "-e", "ip.src",
                    "-e", "ip.dst",
                    "-e", "frame.protocols",
                    "-e", "frame.len",
                    "-e", "_ws.col.Info"
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW
            )
            while self.is_capturing:
                line = self.capture_process.stdout.readline()
                if not line:
                    err = self.capture_process.stderr.read()
                    if err:
                        self.root.after(0, lambda: Messagebox.show_error("tshark error", err))
                    break
                self.process_packet(line.strip())
        except Exception as e:
            Messagebox.show_error("Error", f"Capture failed: {str(e)}")
        finally:
            self.is_capturing = False
            self.root.after(0, lambda: self.start_btn.config(text="Start"))

    def process_packet(self, packet_line):
        """Process a single packet line"""
        fields = packet_line.split("\t")
        if len(fields) != 7:
            return
        packet = {
            "no": fields[0],
            "time": fields[1],
            "src": fields[2],
            "dst": fields[3],
            "protocol": fields[4],
            "length": fields[5],
            "info": fields[6]
        }
        # Detect threats (AI-powered)
        threat, ai_explain, ai_fix = self.detect_threats(packet)
        if threat:
            packet["threat"] = threat
            packet["ai_explain"] = ai_explain
            packet["ai_fix"] = ai_fix
            self.threats.append(packet)
        self.packets.append(packet)
        self.packet_queue.put(packet)

        # Limit packet storage to the most recent MAX_PACKETS packets
        MAX_PACKETS = 10000
        if len(self.packets) > MAX_PACKETS:
            self.packets.pop(0)

    def process_packet_queue(self):
        # Process up to 20 packets per cycle to keep UI responsive
        for _ in range(10):
            if self.packet_queue.empty():
                break
            packet = self.packet_queue.get()
            self.update_packet_display(packet)
        self.root.after(100, self.process_packet_queue)

    def detect_threats(self, packet):
        """
        Use DeepSeek AI to detect threats and get explanation/fix.
        Returns (threat_name, explanation, fix)
        """
        # Simple rules
        if "nmap" in packet["info"].lower():
            return "Port Scan", "Detected Nmap scan", "Block suspicious IP"
        if "malware" in packet["info"].lower():
            return "Malware", "Detected malware signature", "Quarantine the file"
        if "ddos" in packet["info"].lower():
            return "DDoS Attack", "Possible DDoS attack detected", "Rate-limit the IP"
        if re.search(r"(\d{1,3}\.){3}\d{1,3}", packet["info"]):
            return "IP Exposure", "Internal IP address exposed", "Review firewall rules"
        # ...AI detection as fallback...
        prompt = (
            "Analyze the following network packet. If it is a threat, give:\n"
            "- The threat name (e.g., malware, DDoS, port scan, etc.)\n"
            "- A short explanation (max 1-2 sentences)\n"
            "- A short fix (max 1-2 sentences)\n"
            "If not a threat, say 'No threat detected.'\n\n"
            f"Packet Info:\n"
            f"No: {packet['no']}\n"
            f"Time: {packet['time']}\n"
            f"Source: {packet['src']}\n"
            f"Destination: {packet['dst']}\n"
            f"Protocol: {packet['protocol']}\n"
            f"Length: {packet['length']}\n"
            f"Info: {packet['info']}\n"
            "Respond in this format:\n"
            "Threat: <name>\nExplanation: <short explanation>\nFix: <short fix>\n"
        )
        try:
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": self.model_name,
                    "prompt": prompt,
                    "stream": False
                },
                timeout=30
            )
            result = response.json()
            answer = result.get("response", "").strip()
            if "No threat detected" in answer:
                return None, "", ""
            # Parse AI response
            threat = ""
            explain = ""
            fix = ""
            for line in answer.splitlines():
                if line.lower().startswith("threat:"):
                    threat = line.split(":", 1)[1].strip()
                elif line.lower().startswith("explanation:"):
                    explain = line.split(":", 1)[1].strip()
                elif line.lower().startswith("fix:"):
                    fix = line.split(":", 1)[1].strip()
            return threat, explain, fix
        except Exception as e:
            return None, "", ""

    def update_packet_display(self, packet):
        """Update the packet treeview and threat treeview, highlighting threats in both."""
        # Highlight threats in packet_tree
        tags = ()
        src_country = self.get_country_flag(packet["src"])
        dst_country = self.get_country_flag(packet["dst"])
        if "threat" in packet:
            threat_tag = packet["threat"]
            tags = (threat_tag,)
            # Configure the tag color if not already set
            color = self.get_threat_color(threat_tag)
            try:
                self.packet_tree.tag_configure(threat_tag, foreground=color, font=("Segoe UI", 10, "bold"))
            except tb.TclError:
                pass  # Tag may already exist
        self.packet_tree.insert("", tb.END, values=(
            packet["no"],
            packet["time"],
            packet["src"],
            src_country,
            packet["dst"],
            dst_country,
            packet["protocol"],
            packet["length"],
            packet["info"]
        ), tags=tags)

        # If this packet is a threat, add to threat tree as before
        if "threat" in packet:
            color = self.get_threat_color(packet["threat"])
            self.threat_tree.insert("", tb.END, values=(
                packet["no"],
                packet["time"],
                packet["src"],
                src_country,
                packet["dst"],
                dst_country,
                packet["threat"],
                packet["info"],
                packet.get("ai_explain", ""),
                packet.get("ai_fix", "")
            ), tags=(packet["threat"],))
            self.threat_tree.tag_configure(packet["threat"], foreground=color, font=("Segoe UI", 10, "bold"))
            # Play sound only for threats, and throttle (once per second)
            now = time.time()
            if self.sound_var.get() and (now - self.last_sound_time > 1):
                winsound.MessageBeep(winsound.MB_ICONHAND)
                self.last_sound_time = now

        # Auto-scroll if enabled
        if self.autoscroll_var.get():
            self.packet_tree.yview_moveto(1)
            self.threat_tree.yview_moveto(1)
        # Update counts
        self.packet_count.config(text=str(len(self.packets)))
        self.threat_count.config(text=str(len(self.threats)))
        # Update threat analysis if needed
        if "threat" in packet:
            self.update_threat_analysis()
        # Update statistics
        self.update_stats()

    def update_stats(self):
        """Update the statistics chart in the Statistics tab."""
        # Example: plot number of threats over time
        self.stats_data.append(len(self.threats))
        if len(self.stats_data) > 100:
            self.stats_data.pop(0)
        self.ax.clear()
        self.ax.set_title("Threats Over Time")
        self.ax.set_xlabel("Packet #")
        self.ax.set_ylabel("Threat Count")
        self.ax.plot(range(len(self.stats_data)), self.stats_data, color="#FF8800")
        self.stats_canvas.draw()

    def update_threat_analysis(self):
        """Update the threat analysis text and threat treeview"""
        self.threat_tree.delete(*self.threat_tree.get_children())
        if not self.threats:
            return
        show_explain = self.show_ai_explain_var.get()
        show_fix = self.show_ai_fix_var.get()
        for t in self.threats:
            color = self.get_threat_color(t["threat"])
            src_country = self.get_country_flag(t["src"])
            dst_country = self.get_country_flag(t["dst"])
            values = [
                t["no"],
                t["time"],
                t["src"],
                src_country,
                t["dst"],
                dst_country,
                t["threat"],
                t["info"]
            ]
            if show_explain:
                values.append(t.get("ai_explain", ""))
            if show_fix:
                values.append(t.get("ai_fix", ""))
            self.threat_tree.insert("", tb.END, values=values, tags=(t["threat"],))
            self.threat_tree.tag_configure(t["threat"], foreground=color, font=("Segoe UI", 10, "bold"))
        # Update summary as before
        summary = ""
        threat_counts = {}
        for threat in self.threats:
            name = threat["threat"]
            threat_counts[name] = threat_counts.get(name, 0) + 1
        for threat, count in sorted(threat_counts.items(), key=lambda x: -x[1]):
            summary += f"{threat}: {count} detected\n"
        self.threat_summary.config(state="normal")
        self.threat_summary.delete(1.0, tb.END)
        self.threat_summary.insert(tb.END, summary)
        self.threat_summary.config(state="disabled")

    def show_packet_details(self, event):
        item = self.packet_tree.identify_row(event.y)
        if not item:
            selected = self.packet_tree.selection()
            if selected:
                item = selected[0]
            else:
                return
        values = self.packet_tree.item(item, "values")
        if not values or len(values) < 9:
            Messagebox.show_error("Packet Details", "Packet details are incomplete.")
            return
        details = "\n".join([
            f"Packet #{values[0]}",
            f"Time: {values[1]}",
            f"Source: {values[2]} ({values[3]})",
            f"Destination: {values[4]} ({values[5]})",
            f"Protocol: {values[6]}",
            f"Length: {values[7]} bytes",
            f"Info: {values[8]}"
        ])
        Messagebox.show_info(details, "Packet Details")

    def show_threat_details(self, event):
        item = self.threat_tree.selection()
        if not item:
            return
        values = self.threat_tree.item(item, "values")
        if len(values) < 10:
            Messagebox.show_error("Threat Details", "Threat details are incomplete.")
            return
        details = "\n".join([
            f"Packet #{values[0]}",
            f"Time: {values[1]}",
            f"Source: {values[2]} ({values[3]})",
            f"Destination: {values[4]} ({values[5]})",
            f"Threat: {values[6]}",
            f"Info: {values[7]}",
            f"AI Explanation: {values[8]}",
            f"How to Fix: {values[9]}"
        ])
        Messagebox.show_warning(details, "Threat Details")  # <-- Title, Message

    def stop_capture(self):
        """Stop packet capture"""
        self.is_capturing = False
        if self.capture_process:
            self.capture_process.terminate()
        self.start_btn.config(text="Start")

    def stop_nmap_scan(self):
        if self.nmap_process and self.nmap_process.poll() is None:
            self.nmap_process.terminate()
            self.nmap_output.insert(tb.END, "\nScan stopped by user.\n")
            self.nmap_progress_label.config(text="Scan stopped.")
            self.nmap_progress.pack_forget()
            self.nmap_progress_label.pack_forget()

    def clear_data(self):
        """Clear all captured data"""
        self.packet_tree.delete(*self.packet_tree.get_children())
        self.threat_tree.delete(*self.threat_tree.get_children())
        self.threat_summary.delete(1.0, tb.END)
        self.packets.clear()
        self.threats.clear()
        self.packet_count.config(text="0")
        self.threat_count.config(text="0")
        self.stats_data.clear()
        self.ax.clear()
        self.ax.set_title("Threats Over Time")
        self.ax.set_xlabel("Packet #")
        self.ax.set_ylabel("Threat Count")
        self.stats_canvas.draw()

    def export_data(self):
        """Export data to Word (.docx) file"""
        try:
            from docx import Document
            filename = f"capture_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.get_save_path(filename)
            doc = Document()
            doc.add_heading("Network Capture Report", 0)
            doc.add_paragraph(f"Interface: {self.interface}")
            doc.add_paragraph(f"Filter: {self.filter_entry.get()}")
            doc.add_paragraph(f"Timestamp: {datetime.now().isoformat()}")
            doc.add_heading("Packets", level=1)
            for pkt in self.packets:
                doc.add_paragraph(
                    f"No: {pkt['no']}, Time: {pkt['time']}, Src: {pkt['src']}, Dst: {pkt['dst']}, "
                    f"Proto: {pkt['protocol']}, Len: {pkt['length']}, Info: {pkt['info']}, "
                    f"Threat: {pkt.get('threat', 'None')}, Explain: {pkt.get('ai_explain', '')}, Fix: {pkt.get('ai_fix', '')}"
                )
            doc.save(filepath)
            Messagebox.show_info("Export Complete", f"Data saved to {filepath}")
        except Exception as e:
            Messagebox.show_error("Export Error", str(e))

    def export_threats_csv(self):
        """Export threats data to CSV file"""
        try:
            filename = f"threats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            filepath = self.get_save_path(filename)
            if not filepath:
                return  # User cancelled
            with open(filepath, "w", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(["No", "Time", "Source", "Destination", "Threat", "Info", "AI Explain", "Fix"])
                for t in self.threats:
                    writer.writerow([
                        t["no"], t["time"], t["src"], t["dst"], t["threat"], t["info"], t.get("ai_explain", ""), t.get("ai_fix", "")
                    ])
            Messagebox.show_info(f"Threats saved to {filepath}", "Export Complete")
        except Exception as e:
            Messagebox.show_error(str(e), "Export Error")

    def export_threats_pdf(self):
        """Export threats data to PDF file"""
        filename = f"threats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = self.get_save_path(filename)
        if not filepath:
            return  # User cancelled
        c = canvas.Canvas(filepath, pagesize=letter)
        c.drawString(100, 750, "Threat Report")
        y = 700
        for t in self.threats:
            c.drawString(50, y, f"{t['no']} {t['time']} {t['src']} {t['dst']} {t['threat']} {t['info']} {t.get('ai_explain', '')} {t.get('ai_fix', '')}")
            y -= 20
            if y < 50:
                c.showPage()
                y = 750
        c.save()
        Messagebox.show_info("Export Complete", f"Threats saved to {filepath}")

    def create_ai_report(self):
        """Generate a detailed AI-written Word report for all threats."""
        if not self.threats:
            Messagebox.show_info("No Threats", "No threats detected to report.")
            return

        def report_thread():
            logs = []
            for t in self.threats:
                src_country = self.get_country_flag(t["src"])
                dst_country = self.get_country_flag(t["dst"])
                logs.append(
                    f"No: {t['no']}, Time: {t['time']}, Src: {t['src']} ({src_country}), Dst: {t['dst']} ({dst_country}), "
                    f"Proto: {t['protocol']}, Len: {t['length']}, Info: {t['info']}, "
                    f"Threat: {t.get('threat','')}, Explain: {t.get('ai_explain','')}, Fix: {t.get('ai_fix','')}"
                )
            log_content = "\n".join(logs)
            print("DEBUG: Threat logs sent to AI:\n", log_content)  # For debugging

            prompt = (
                "You are a cybersecurity expert. Analyze ALL the following detected network threats and provide a structured, detailed report. "
                "For each threat, output in this format:\n"
                "1. Threat Number: <No>\n"
                "   Name: <Threat Name>\n"
                "   Degree of Risk: (Low/Medium/High/Critical) (based on context and threat type)\n"
                "   Source: <Src>   Destination: <Dst>\n"
                "   Time: <Time>\n"
                "   Protocol: <Proto>\n"
                "   Length: <Len>\n"
                "   Info: <Info>\n"
                "   Explanation: <Short explanation>\n"
                "   Fix: <Short fix>\n"
                "List all threats in order, and at the end, provide a summary of the overall risk and your recommendations. "
                "Be specific, do not include disclaimers or meta-comments, and do not mention being an AI or model. "
                "Only output the report body, with no introduction or closing remarks.\n\n"
                f"Threat Logs:\n{log_content}\n\n"
                "Structured Threat Analysis:"
            )

            try:
                response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": self.model_name,
                        "prompt": prompt,
                        "stream": False
                    },
                    timeout=60
                )
                result = response.json()
                report = result.get("response", "").strip()
            except Exception as e:
                self.root.after(0, lambda: Messagebox.show_error("Report Error", str(e)))
                return

            from docx import Document
            from datetime import datetime
            report_filename = f"AI_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_filepath = self.get_save_path(report_filename)
            if not report_filepath:
                return
            doc = Document()
            doc.add_heading("Network Threat Report", 0)
            for line in report.splitlines():
                doc.add_paragraph(line)
            doc.save(report_filepath)
            self.root.after(0, lambda: Messagebox.show_info("Report Generated", f"AI report saved to {report_filepath}"))

        threading.Thread(target=report_thread, daemon=True).start()

    def save_report_to_pdf(self, report_content, filepath):
        """Save the AI-generated report content to a PDF file."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter

        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, height - 72, "Network Threat Report")

        # Content
        c.setFont("Helvetica", 12)
        text_object = c.beginText(72, height - 100)
        text_object.setTextOrigin(72, height - 100)
        text_object.setLeading(14)
        for line in report_content.splitlines():
            text_object.textLine(line)
        c.drawText(text_object)

        c.save()

    def get_save_path(self, default_filename):
        """Open a file dialog to get the save path from the user"""
        
        filepath = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("CSV files", "*.csv"), ("PDF files", "*.pdf")],
            title="Save As",
            initialfile=default_filename
        )
        return filepath

    def toggle_dark_mode(self):
        style = tb.Style()
        if self.dark_mode:
            style.theme_use('flatly')  # Light theme
        else:
            style.theme_use('cyborg')  # Dark theme
        self.dark_mode = not self.dark_mode

    def save_settings(self):
        """Save settings such as AI model name and user options to a file"""
        try:
            settings = {
                "model_name": self.model_entry.get(),
                "sound": self.sound_var.get(),
                "autoscroll": self.autoscroll_var.get(),
                "darkmode": self.darkmode_var.get(),
                "autosave": self.autosave_var.get(),
                "show_tooltips": self.show_tooltips_var.get(),
                "confirm_stop": self.confirm_stop_var.get(),
                "save_logs": self.save_logs_var.get(),
                "show_ai_explain": self.show_ai_explain_var.get(),
                "show_ai_fix": self.show_ai_fix_var.get()
            }
            with open("settings.json", "w") as f:
                json.dump(settings, f, indent=2)
            Messagebox.show_info("Settings", "Settings saved successfully")
        except Exception as e:
            Messagebox.show_error("Settings Error", str(e))

    def load_settings(self):
        """Load settings from a file"""
        try:
            if os.path.exists("settings.json"):
                with open("settings.json", "r") as f:
                    settings = json.load(f)
                    self.model_name = settings.get("model_name", "deepseek-coder")
                    self.model_entry.delete(0, tb.END)
                    self.model_entry.insert(0, self.model_name)
                    self.sound_var.set(settings.get("sound", True))
                    self.autoscroll_var.set(settings.get("autoscroll", True))
                    self.darkmode_var.set(settings.get("darkmode", False))
                    self.autosave_var.set(settings.get("autosave", False))
                    self.show_tooltips_var.set(settings.get("show_tooltips", True))
                    self.confirm_stop_var.set(settings.get("confirm_stop", True))
                    self.save_logs_var.set(settings.get("save_logs", True))
                    self.show_ai_explain_var.set(settings.get("show_ai_explain", True))
                    self.show_ai_fix_var.set(settings.get("show_ai_fix", True))
        except Exception as e:
            Messagebox.show_error("Settings Error", str(e))

    def run_nmap_scan(self):
        target = self.nmap_target_entry.get().strip()
        if not self.is_nmap_installed():
            Messagebox.show_error("Nmap Error", "Nmap is not installed or not in your PATH. Please install Nmap from htbs://nmap.org/download.html and ensure it is accessible.")
            return
        if not target or not self.is_valid_target(target):
            Messagebox.show_error("Nmap", "Please enter a valid IP address or domain.")
            return

        nmap_path = self.get_nmap_path()
        nmap_args = [nmap_path]

        # IPv6
        if self.nmap_ipv6_var.get():
            nmap_args.append("-6")
        # No ping
        if self.nmap_no_ping_var.get():
            nmap_args.append("-Pn")
        # Verbose/debug
        if self.nmap_verbose_var.get():
            nmap_args.append("-v")
        if self.nmap_debug_var.get():
            nmap_args.append("-d")
        # Timing
        timing = self.nmap_timing_combo.get()
        if timing:
            nmap_args.append(f"-{timing}")
        # Scan type
        scan_type = self.nmap_scan_type.get()
        if scan_type == "Fast":
            nmap_args += ["-F"]
            scan_timeout = 60
        elif scan_type == "Medium":
            nmap_args += ["-sV", "-T4", "--top-ports", "100", "--script", "vuln"]
            scan_timeout = 300
        elif scan_type == "Full":
            if not Messagebox.yesno("Warning", "Full scans may impact network performance. Continue?"):
                return
            nmap_args += ["-A", "-p-", "--script", "vuln"]
            scan_timeout = 600
        else:  # Custom
            scan_timeout = 300

        # Custom ports
        ports = self.nmap_ports_entry.get().strip()
        if ports and "-p" not in nmap_args:
            nmap_args += ["-p", ports]
        # Aggressive, OS, Service
        if self.nmap_aggressive_var.get() and "-A" not in nmap_args:
            nmap_args.append("-A")
        if self.nmap_osdetect_var.get() and "-O" not in nmap_args:
            nmap_args.append("-O")
        if self.nmap_service_var.get() and "-sV" not in nmap_args:
            nmap_args.append("-sV")
        # NSE scripts
        scripts = self.nmap_scripts_entry.get().strip()
        if scripts:
            nmap_args += ["--script", scripts]
        # Output format
        out_fmt = self.nmap_output_format.get()
        if out_fmt == "XML":
            nmap_args += ["-oX", "nmap_output.xml"]
        elif out_fmt == "Grepable":
            nmap_args += ["-oG", "nmap_output.grep"]
        # Extra arguments
        extra = self.nmap_extra_entry.get().strip()
        if extra:
            if "-p" in extra:
                Messagebox.show_error("Nmap", "Please specify custom ports only in the Ports field, not in Extra Args.")
                return
            nmap_args += extra.split()
        nmap_args.append(target)

        self.nmap_output.config(state="normal")
        self.nmap_output.delete(1.0, tb.END)  # Clear previous output
        self.nmap_output.insert(tb.END, f"Running {self.nmap_scan_type.get()} nmap scan on {target}...\n")
        self.nmap_output.update()
        self.nmap_progress["value"] = 0
        self.nmap_progress["maximum"] = 100
        self.nmap_progress.pack(fill=tb.X, pady=(5, 0))
        self.nmap_progress_label.config(text="Starting scan...")
        self.nmap_progress_label.pack(fill=tb.X)
        self.nmap_progress_running = True

        def nmap_thread():
            try:
                self.nmap_process = subprocess.Popen(
                    nmap_args,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    text=True,
                    bufsize=1
                )
                current_port = ""
                for line in self.nmap_process.stdout:
                    self.root.after(0, lambda l=line: self.nmap_output.insert(tb.END, l))
                    port_match = re.search(r"(\d{1,5})/(tcp|udp)\s+open", line)
                    if port_match:
                        current_port = port_match.group(1)
                        self.root.after(0, lambda p=current_port: self.nmap_progress_label.config(text=f"Scanning port: {p}"))
                    elif "Nmap scan report for" in line:
                        self.root.after(0, lambda: self.nmap_progress_label.config(text="Host detected, scanning ports..."))
                self.nmap_process.wait()
                self.nmap_progress_running = False
                self.root.after(0, self.nmap_progress.pack_forget)
                self.root.after(0, self.nmap_progress_label.pack_forget)
                self.root.after(0, lambda: self.nmap_stop_btn.config(state="disabled"))
                output = self.nmap_output.get(1.0, tb.END)
                geo_output = self.annotate_nmap_output_with_geo(output)
                self.root.after(0, lambda: self.nmap_output.insert(tb.END, geo_output))
                self.root.after(0, lambda: self.show_nmap_vulns(geo_output))
            except Exception as e:
                self.root.after(0, lambda err=e: self.nmap_output.insert(tb.END, f"\nError running nmap: {err}\n"))
                self.root.after(0, lambda: self.nmap_stop_btn.config(state="disabled"))
        threading.Thread(target=nmap_thread, daemon=True).start()

    def show_nmap_vulns(self, nmap_output):
        self.nmap_output.config(state="normal")
        vulns = []
        open_ports = []
        for line in nmap_output.splitlines():
            if "open" in line and "/" in line:
                open_ports.append(line)
            if "VULNERABLE" in line or "CVE-" in line:
                vulns.append(line)
        # Tag configuration for colors
        self.nmap_output.tag_config("header", foreground="#FF8800", font=("Segoe UI", 12, "bold"))
        self.nmap_output.tag_config("port", foreground="#2979FF", font=("Segoe UI", 11, "bold"))
        self.nmap_output.tag_config("vuln", foreground="#FF1744", font=("Segoe UI", 11, "bold"))
        self.nmap_output.tag_config("ai", foreground="#00C853", font=("Segoe UI", 11, "italic"))
        self.nmap_output.tag_config("normal", foreground="#FFFFFF", font=("Segoe UI", 11))
        # Insert open ports
        if open_ports:
            self.nmap_output.insert(tb.END, "\n\nOpen Ports:\n", "header")
            for port in open_ports:
                self.nmap_output.insert(tb.END, port + "\n", "port")
        # Insert vulnerabilities
        if vulns:
            self.nmap_output.insert(tb.END, "\n\nDetected Vulnerabilities:\n", "header")
            for v in vulns:
                self.nmap_output.insert(tb.END, v + "\n", "vuln")
        else:
            self.nmap_output.insert(tb.END, "\n\nNo explicit vulnerabilities found by Nmap scripts.\n", "normal")
        self.nmap_output.config(state="disabled")

    def analyze_nmap_with_ai(self):
        nmap_text = self.nmap_output.get(1.0, tb.END).strip()
        if not nmap_text or "0 hosts up" in nmap_text or "Nmap done" in nmap_text and "open" not in nmap_text:
            Messagebox.show_info("Nmap", "No useful Nmap results to analyze. Try a different target.")
            return
        self.nmap_output.config(state="normal")
        self.nmap_output.insert(tb.END, "\n\nAI is analyzing the scan...\n", "header")
        self.nmap_output.config(state="disabled")
        self.nmap_output.update()
        def ai_thread():
            prompt = (
                "Analyze the following Nmap scan result. Identify any open ports, detected vulnerabilities, "
                "and security risks. For each vulnerability or risk, explain what it means and suggest fixes. "
                "If only standard web ports (80/443) are open and no vulnerabilities are found, reply with: "
                "'No vulnerabilities detected. Only standard web ports (80/443) are open. Ensure your web server and SSL certificates are up to date.'\n\n"
                f"Nmap Output:\n{nmap_text}\n"
                "AI Analysis:"
            )
            try:
                response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": self.model_name,
                        "prompt": prompt,
                        "stream": False
                    },
                    timeout=120
                )
                result = response.json()
                ai_msg = result.get("response", "").strip()
            except Exception as e:
                ai_msg = f"DeepSeek error: {e}"
            # Show in Nmap output
            def show_ai():
                self.nmap_output.config(state="normal")
                self.nmap_output.insert(tb.END, "\n\nAI Analysis:\n", "header")
                self.nmap_output.insert(tb.END, ai_msg + "\n", "ai")
                self.nmap_output.config(state="disabled")
            self.root.after(0, show_ai)
            # Show in chat box, formatted and organized
            def show_in_chat():
                self.chat_display.config(state="normal")
                self.chat_display.insert(tb.END, "\n" + "="*40 + "\n", "ai")
                self.chat_display.insert(tb.END, "AI (Nmap Analysis):\n", "ai")
                self.chat_display.insert(tb.END, ai_msg + "\n", "ai")
                self.chat_display.insert(tb.END, "="*40 + "\n\n", "ai")
                self.chat_display.config(state="disabled")
                self.chat_display.see(tb.END)
            self.root.after(0, show_in_chat)
        threading.Thread(target=ai_thread, daemon=True).start()

    def analyze_threats_with_ai(self):
        """Send all detected threats to the AI and show the analysis in both the threats tab and chat box."""
        if not self.threats:
            Messagebox.show_info("AI Threat Analysis", "No threats detected to analyze.")
            return

        # Prepare the logs for AI
        logs = []
        for t in self.threats:
            src_country = self.get_country_flag(t["src"])
            dst_country = self.get_country_flag(t["dst"])
            logs.append(
                f"No: {t['no']}, Time: {t['time']}, Src: {t['src']} ({src_country}), Dst: {t['dst']} ({dst_country}), "
                f"Proto: {t['protocol']}, Len: {t['length']}, Info: {t['info']}, "
                f"Threat: {t.get('threat','')}, Explain: {t.get('ai_explain','')}, Fix: {t.get('ai_fix','')}"
            )
        log_content = "\n".join(logs)

        prompt = (
            "You are a cybersecurity expert. Analyze ALL the following detected network threats and provide a structured, detailed report. "
            "For each threat, output in this format:\n"
            "1. Threat Number: <No>\n"
            "   Name: <Threat Name>\n"
            "   Degree of Risk: (Low/Medium/High/Critical) (based on context and threat type)\n"
            "   Source: <Src>   Destination: <Dst>\n"
            "   Time: <Time>\n"
            "   Protocol: <Proto>\n"
            "   Length: <Len>\n"
            "   Info: <Info>\n"
            "   Explanation: <Short explanation>\n"
            "   Fix: <Short fix>\n"
            "List all threats in order, and at the end, provide a summary of the overall risk and your recommendations. "
            "Be specific, do not include disclaimers or meta-comments, and do not mention being an AI or model. "
            "Only output the report body, with no introduction or closing remarks.\n\n"
            f"Threat Logs:\n{log_content}\n\n"
            "Structured Threat Analysis:"
        )

        def ai_thread():
            try:
                response = requests.post( 
                    "http://localhost:11434/api/generate",
                    json={
                        "model": self.model_name,
                        "prompt": prompt,
                        "stream": False
                    },
                    timeout=120
                )
                result = response.json()
                ai_msg = result.get("response", "").strip()
            except Exception as e:
                ai_msg = f"DeepSeek error: {e}"

            # Show in a popup and in chat
            self.root.after(0, lambda: Messagebox.show_info(ai_msg, "AI Threat Analysis"))
            def show_in_chat():
                self.chat_display.config(state="normal")
                self.chat_display.insert(tb.END, "\n" + "="*40 + "\n", "ai")
                self.chat_display.insert(tb.END, "AI (Threats Analysis):\n", "ai")
                self.chat_display.insert(tb.END, ai_msg + "\n", "ai")
                self.chat_display.insert(tb.END, "="*40 + "\n\n", "ai")
                self.chat_display.config(state="disabled")
                self.chat_display.see(tb.END)
            self.root.after(0, show_in_chat)
        threading.Thread(target=ai_thread, daemon=True).start()

    def create_nmap_report(self):
        """Generate a detailed AI-written Word report for the last Nmap scan."""
        nmap_text = self.nmap_output.get(1.0, tb.END).strip()
        if not nmap_text or "Nmap done" not in nmap_text:
            Messagebox.show_info("Nmap Report", "No Nmap scan results to report.")
            return

        def report_thread():
            prompt = (
                "Write a professional, concise network scan report based only on the following Nmap output. "
                "For each open port or detected vulnerability, explain what it is, why it is important, and how to fix it. "
                "Summarize the overall risk and give recommendations. "
                "Do NOT include any disclaimers, meta-comments, or information about yourself. "
                "Do NOT mention being an AI, model, or assistant. "
                "Only output the report body, with no introduction or closing remarks.\n\n"
                f"Nmap Output:\n{nmap_text}\n\n"
                "Report:"
            )
            try:
                response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": self.model_name,
                        "prompt": prompt,
                        "stream": False
                    },
                    timeout=120
                )
                result = response.json()
                report = result.get("response", "").strip()
            except Exception as e:
                self.root.after(0, lambda: Messagebox.show_error("Nmap Report Error", str(e)))
                return

            from docx import Document
            from datetime import datetime
            report_filename = f"Nmap_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_filepath = self.get_save_path(report_filename)
            if not report_filepath:
                return
            doc = Document()
            doc.add_heading("Nmap Scan Report", 0)
            for line in report.splitlines():
                doc.add_paragraph(line)
            doc.save(report_filepath)
            self.root.after(0, lambda: Messagebox.show_info("Nmap Report Generated", f"AI Nmap report saved to {report_filepath}"))

        threading.Thread(target=report_thread, daemon=True).start()

    def export_packets_prompt(self):
        """Prompt user to choose export format for packets."""
        choice = Messagebox.yesno("Export Format", "Export packets as JSON?\n(Click 'No' for CSV)")
        if choice:
            self.export_packets_json()
        else:
            self.export_packets_csv()

    def export_packets_json(self):
        """Export all packets to a JSON file."""
        filename = f"packets_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        filepath = self.get_save_path(filename)
        if not filepath:
            return
        try:
            with open(filepath, "w", newline="", encoding="utf-8") as f:
                json.dump(self.packets, f, indent=2)
            Messagebox.show_info(f"Data saved to {filepath}", "Export Complete")
        except Exception as e:
            Messagebox.show_error(str(e), "Export Error")

    def export_packets_csv(self):
        """Export all packets to a CSV file, marking threats and adding a threat color column."""
        filename = f"packets_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        filepath = self.get_save_path(filename)
        if not filepath:
            return
        try:
            with open(filepath, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow([
                    "No", "Time", "Source", "Destination", "Protocol", "Length", "Info", "Threat", "AI Explain", "Fix", "Threat Color"
                ])
                for pkt in self.packets:
                    threat = pkt.get("threat", "")
                    color = self.get_threat_color(threat) if threat else ""
                    writer.writerow([
                        pkt["no"], pkt["time"], pkt["src"], pkt["dst"], pkt["protocol"], pkt["length"], pkt["info"],
                        threat, pkt.get("ai_explain", ""), pkt.get("ai_fix", ""), color
                    ])
            Messagebox.show_info(f"Packets saved to {filepath}", "Export Complete")
        except Exception as e:
            Messagebox.show_error(str(e), "Export Error")

    def show_nmap_version(self):
        try:
            nmap_path = self.get_nmap_path()
            result = subprocess.run([nmap_path, "--version"], capture_output=True, text=True, timeout=10)
            Messagebox.show_info("Nmap Version", result.stdout)
        except Exception as e:
            Messagebox.show_error("Nmap Version Error", str(e))

    def save_nmap_profile(self):
        """Save current Nmap scan settings to a JSON file."""
        profile = {
            "target": self.nmap_target_entry.get(),
            "scan_type": self.nmap_scan_type.get(),
            "ports": self.nmap_ports_entry.get(),
            "aggressive": self.nmap_aggressive_var.get(),
            "osdetect": self.nmap_osdetect_var.get(),
            "service": self.nmap_service_var.get(),
            "timing": self.nmap_timing_combo.get(),
            "scripts": self.nmap_scripts_entry.get(),
            "ipv6": self.nmap_ipv6_var.get(),
            "no_ping": self.nmap_no_ping_var.get(),
            "verbose": self.nmap_verbose_var.get(),
            "debug": self.nmap_debug_var.get(),
            "output_format": self.nmap_output_format.get(),
            "extra": self.nmap_extra_entry.get()
        }
        filepath = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON files", "*.json")], title="Save Nmap Profile")
        if filepath:
            with open(filepath, "w") as f:
                json.dump(profile, f, indent=2)
            Messagebox.show_info(f"Profile saved to {filepath}", "Profile Saved")

    def load_nmap_profile(self):
        """Load Nmap scan settings from a JSON file."""
        filepath = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")], title="Load Nmap Profile")
        if filepath and os.path.exists(filepath):
            with open(filepath, "r") as f:
                profile = json.load(f)
            self.nmap_target_entry.delete(0, tb.END)
            self.nmap_target_entry.insert(0, profile.get("target", ""))
            self.nmap_scan_type.set(profile.get("scan_type", "Fast"))
            self.nmap_ports_entry.delete(0, tb.END)
            self.nmap_ports_entry.insert(0, profile.get("ports", ""))
            self.nmap_aggressive_var.set(profile.get("aggressive", False))
            self.nmap_osdetect_var.set(profile.get("osdetect", False))
            self.nmap_service_var.set(profile.get("service", False))
            self.nmap_timing_combo.set(profile.get("timing", "T4"))
            self.nmap_scripts_entry.delete(0, tb.END)
            self.nmap_scripts_entry.insert(0, profile.get("scripts", ""))
            self.nmap_ipv6_var.set(profile.get("ipv6", False))
            self.nmap_no_ping_var.set(profile.get("no_ping", False))
            self.nmap_verbose_var.set(profile.get("verbose", False))
            self.nmap_debug_var.set(profile.get("debug", False))
            self.nmap_output_format.set(profile.get("output_format", "Normal"))
            self.nmap_extra_entry.delete(0, tb.END)
            self.nmap_extra_entry.insert(0, profile.get("extra", ""))
            Messagebox.show_info("Profile Loaded", f"Profile loaded from {filepath}")

    def annotate_nmap_output_with_geo(self, nmap_output):
        """
        Optionally annotate Nmap output with country/flag info for each IP.
        For now, just return the output unchanged.
        """
        return nmap_output

    def search_packets(self):
        """Filter packets in the packet_tree by search text."""
        search_text = self.search_entry.get().strip().lower()
        # Remove all current rows
        for item in self.packet_tree.get_children():
            self.packet_tree.delete(item)
        # Re-insert only matching packets
        for pkt in self.packets:
            row = (
                pkt["no"],
                pkt["time"],
                pkt["src"],
                self.get_country_flag(pkt["src"]),
                pkt["dst"],
                self.get_country_flag(pkt["dst"]),
                pkt["protocol"],
                pkt["length"],
                pkt["info"]
            )
            if any(search_text in str(field).lower() for field in row):
                self.packet_tree.insert("", tb.END, values=row)

    def send_chat_message(self):
        """Handle sending a chat message in the AI Chat tab."""
        user_msg = self.chat_entry.get().strip()
        if not user_msg:
            return
        self.chat_display.config(state="normal")
        self.chat_display.insert(tb.END, f"You: {user_msg}\n", "user")
        self.chat_display.config(state="disabled")
        self.chat_display.see(tb.END)
        self.chat_entry.delete(0, tb.END)

        # Show "thinking..." while waiting for AI
        self.chat_display.config(state="normal")
        self.chat_display.insert(tb.END, "AI is thinking...\n", "thinking")
        self.chat_display.config(state="disabled")
        self.chat_display.see(tb.END)

        def ai_thread():
            prompt = (
                "You are a cybersecurity assistant. Answer the following user question or request in a helpful, concise way.\n"
                "User: " + user_msg
            )
            try:
                response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": self.model_name,
                        "prompt": prompt,
                        "stream": False
                    },
                    timeout=60
                )
                result = response.json()
                ai_msg = result.get("response", "").strip()
            except Exception as e:
                ai_msg = f"AI error: {e}"

            def show_ai_response():
                # Remove the "thinking..." line
                self.chat_display.config(state="normal")
                self.chat_display.delete("end-2l", "end-1l")
                self.chat_display.insert(tb.END, f"AI: {ai_msg}\n", "ai")
                self.chat_display.config(state="disabled")
                self.chat_display.see(tb.END)

            self.root.after(0, show_ai_response)

        threading.Thread(target=ai_thread, daemon=True).start()

    def on_closing(self):
       
        """Handle cleanup and close the application safely."""
        # Stop any running capture
        self.is_capturing = False
        if hasattr(self, "capture_process") and self.capture_process:
            try:
                self.capture_process.terminate()
            except Exception:
                pass
        # Close GeoIP reader if open
        if hasattr(self, "geoip_reader") and self.geoip_reader:
            try:
                self.geoip_reader.close()
            except Exception:
                pass
        self.root.destroy()

    def get_threat_color(self, threat):
        """Return a color code for a given threat type."""
        threat = (threat or "").lower()
        if "malware" in threat:
            return "#FF1744"  # Red
        if "ddos" in threat:
            return "#FF9100"  # Orange
        if "port scan" in threat or "scan" in threat:
            return "#2979FF"  # Blue
        if "ip exposure" in threat or "exposure" in threat:
            return "#D500F9"  # Purple
        if "error" in threat or "fail" in threat:
            return "#C51162"  # Pink
        return "#222222"      # Default (dark gray)

    def is_nmap_installed(self):
        """Check if Nmap is installed and accessible."""
        try:
            result = subprocess.run(["nmap", "--version"], capture_output=True, text=True, timeout=5)
            return result.returncode == 0
        except Exception:
            return False

    def is_valid_ip(self, ip):
        try:
            socket.inet_aton(ip)
            return True
        except socket.error:
            return False

    def add_tooltip(widget, text):
        tooltip = tb.Toplevel(widget)
        tooltip.withdraw()
        tooltip.overrideredirect(True)
        label = tb.Label(tooltip, text=text, background="#ffffe0", relief="solid", borderwidth=1)
        label.pack()
        def enter(event): tooltip.deiconify()
        def leave(event): tooltip.withdraw()
        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)

    def is_valid_target(self, target):
        """Check if the target is a valid IP address or hostname."""
        if not target:
            return False
        # Simple check: no spaces, and at least one dot for hostname or IP
        if " " in target or "." not in target:
                                                                                         return False

        # If it looks like an IP, check if it's valid
        if re.match(r"^\d{1,3}(\.\d{1,3}){3}$", target):
            return self.is_valid_ip(target)
        # Otherwise, it should be a hostname (more checks can be added)
        return True
    
    def get_nmap_path(self):
        """Return the path to the nmap executable."""
        # Adjust this path if nmap is installed elsewhere
        possible_paths = [
            r"C:\Program Files (x86)\Nmap\nmap.exe",
            r"C:\Program Files\Nmap\nmap.exe",
            "nmap"  # fallback to PATH
        ]
        for path in possible_paths:
            if os.path.exists(path) or path == "nmap":
                return path
        # If not found, fallback to just "nmap" (will fail if not in PATH)
        return "nmap"
    
    # --- Pentest Tools Tab ---
    def run_ping(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target.", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nPing {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def ping_thread():
            cmd = ["ping", "-n", "4", target] if platform.system() == "Windows" else ["ping", "-c", "4", target]
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
                output = result.stdout
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=ping_thread, daemon=True).start()

    def run_traceroute(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target.", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nTraceroute {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def trace_thread():
            cmd = ["tracert", target] if platform.system() == "Windows" else ["traceroute", target]
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                output = result.stdout
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=trace_thread, daemon=True).start()

    def run_whois(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target.", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nWhois {target}:\n", "header")
        self.pentest_output.config(state="disabled")

        def whois_thread():
            try:
                if re.match(r"^\d{1,3}(\.\d{1,3}){3}$", target):
                    # Use ipinfo.io for IP whois
                    resp = requests.get(f"https://ipinfo.io/{target}/json", timeout=10)
                    if resp.status_code == 200:
                        output = resp.text
                    else:
                        output = f"Error: Could not fetch info for {target}"
                else:
                    import whois
                    w = whois.whois(target)
                    output = str(w)
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))

        threading.Thread(target=whois_thread, daemon=True).start()

    def run_nslookup(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target.", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nNslookup {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def nslookup_thread():
            cmd = ["nslookup", target]
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
                output = result.stdout
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=nslookup_thread, daemon=True).start()

    def run_portscan(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target.", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nPort Scan (Nmap) {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def portscan_thread():
            cmd = [self.get_nmap_path(), "-F", target]
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                output = result.stdout
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=portscan_thread, daemon=True).start()

    def run_http_headers(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target (domain or URL).", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nHTTP Headers for {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def headers_thread():
            cmd = ["curl", "-I", target]
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
                output = result.stdout
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=headers_thread, daemon=True).start()

    def run_ssl_info(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a target (domain).", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nSSL Info for {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def ssl_thread():
            cmd = ["openssl", "s_client", "-connect", f"{target}:443", "-servername", target]
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
                output = result.stdout
            except Exception as e:
                output = f"Error: {e}"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=ssl_thread, daemon=True).start()

    def run_dns_zone_transfer(self):
        target = self.pentest_target_entry.get().strip()
        if not target:
            Messagebox.show_error("Please enter a domain.", "Pentest")
            return
        self.pentest_output.config(state="normal")
        self.pentest_output.delete(1.0, tb.END)
        self.pentest_output.insert(tb.END, f"\n\nDNS Zone Transfer Test for {target}:\n", "header")
        self.pentest_output.config(state="disabled")
        def zone_thread():
            # Try common nameservers
            nameservers = ["ns1." + target, "ns2." + target]
            output = ""
            for ns in nameservers:
                cmd = ["nslookup", "-type=ns", target]
                try:
                    ns_result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                    output += f"\nNameservers for {target}:\n{ns_result.stdout}\n"
                    # Try zone transfer
                    cmd2 = ["nslookup", "-type=any", target, ns]
                    zt_result = subprocess.run(cmd2, capture_output=True, text=True, timeout=15)
                    output += f"\nZone transfer attempt with {ns}:\n{zt_result.stdout}\n"
                except Exception as e:
                    output += f"Error: {e}\n"
            self.root.after(0, lambda: self.display_pentest_output(output))
        threading.Thread(target=zone_thread, daemon=True).start()

    def display_pentest_output(self, output):
        self.pentest_output.config(state="normal")
        
        self.pentest_output.insert(tb.END, output + "\n", "output")
        self.pentest_output.config(state="disabled")
        self.pentest_output.see(tb.END)

if __name__ == "__main__":
    root = tb.Window(themename="cyborg")  # or "cyborg" for dark mode
    app = WireSharkAnalyzer(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()